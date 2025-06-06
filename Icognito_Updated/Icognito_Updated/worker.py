import requests
from bs4 import BeautifulSoup
import csv
from docx import Document
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import boto3
from langchain_community.embeddings import BedrockEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
import requests
from urllib.parse import urljoin, urlparse
import textwrap
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from threading import Lock
import asyncio
import logging

# Global lock for thread-safe operations
lock = Lock()

# Basic configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("main.log"),        # Log to a file
        logging.StreamHandler()                    # Also log to console
    ]
)

logger = logging.getLogger(__name__)

# Define headers, including a common browser's User-Agent string
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
}

# Function to clean up text
def clean_text(text):
    return ' '.join(text.split()).strip()

# Step 1: Read all href URLs from a given website and create a CSV
def extract_urls_to_csv(website_url, folder_name):
    response = requests.get(website_url)
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Find all href links
    urls = [a['href'] for a in soup.find_all('a', href=True)]
    
    # Write URLs to CSV
    csv_file_name = os.path.join(folder_name, 'urls.csv')
    with open(csv_file_name, mode='w', newline='') as file:
        writer = csv.writer(file)
        for url in urls:
            writer.writerow([url])
    
    print(f"Extracted {len(urls)} URLs and saved to {csv_file_name}")
    return urls

# Step 2: Read URLs from CSV and extract specified tags into separate .docx files
def extract_content_from_urls(urls, folder_name):
    for url in urls:
        try:
            response = requests.get(url,verify=False)
            print(response)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Create a new document for each URL
            document = Document()
            
            # Extract h1, h2, p, div, span tags
            for tag in ['h1', 'h2', 'p', 'span']:
                elements = soup.find_all(tag)
                for element in elements:
                    cleaned_text = clean_text(element.get_text())
                    if tag in ['h1', 'h2']:
                        # Add header in bold
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run(cleaned_text)
                        run.bold = True
                    else:
                        # Add normal text
                        document.add_paragraph(cleaned_text)
            
            # Create a filename based on the URL
            safe_url = url.replace('https://', '').replace('http://', '').replace('/', '_')[:50]  # Shorten if necessary
            docx_file_name = os.path.join(folder_name, f"{safe_url}.docx")
            
            document.save(docx_file_name)
            print(f"Content extracted and saved to {docx_file_name}")
            
        except Exception as e:
            print(f"Failed to fetch {url}: {e}")


# Function to create directory to store files
def create_directory(directory):
    logger.info(f"Creating directory: {directory}")
    if not os.path.exists(directory):
        os.makedirs(directory)

# Function to scrape the website vishwanth change to parallel processing
# def scrape_website(base_url):
#     visited = set()
#     pages_to_visit = [base_url]
    
#     # Create a directory to store the scraped pages
#     domain_name = urlparse(base_url).netloc.replace("www.", "")
#     create_directory(domain_name)

#     while pages_to_visit:
#         url = pages_to_visit.pop(0)
#         print(f"Visiting: {url}")

#         if url in visited:
#             continue

#         visited.add(url)

#         try:
#             response = requests.get(url,headers=headers)
#             response.raise_for_status()  # Check for HTTP errors

#             # Parse the HTML content
#             soup = BeautifulSoup(response.text, "html.parser")
            
#             # Create a new Document
#             doc = Document()

#             # Scrape and add h, p, and span tags to the document
#             for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'span']):
#                 if tag.name.startswith('h'):
#                     doc.add_heading(tag.get_text(strip=True), level=int(tag.name[-1]))
#                 elif tag.name == 'p':
#                     doc.add_paragraph(tag.get_text(strip=True))
#                 elif tag.name == 'span':
#                     doc.add_paragraph(tag.get_text(strip=True), style='Normal')

#             # Save the document as a .docx file
#             page_name = urlparse(url).path.strip("/").replace("/", "_") or "index"
#             doc_file_path = os.path.join(domain_name, f"{page_name}.docx")

#             doc.save(doc_file_path)
                
#             print(f"Saved: {doc_file_path}")

#             # Find all links on the page
#             for link in soup.find_all("a"):
#                 href = link.get("href")
#                 if href:
#                     full_url = urljoin(url, href)
                    
#                     # Exclude image URLs
#                     if not (full_url.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg','.pdf','.docx'))):
#                         # Add the link if it is part of the same domain and not visited
#                         if urlparse(full_url).netloc.replace("www.", "") == domain_name and full_url not in visited:
#                             pages_to_visit.append(full_url)

#         except Exception as e:
#             print(f"Failed to process {url}: {e}")

#multithreaded version of scrape_website
# def scrape_single_page(url, domain_name, visited, pages_to_visit, excluded_extensions, headers):
#     with lock:
#         if url in visited or any(url.lower().endswith(ext) for ext in excluded_extensions):
#             return
#         visited.add(url)

#     print(f"Visiting: {url}")
#     try:
#         response = requests.get(url, headers=headers, timeout=10)
#         response.raise_for_status()

#         soup = BeautifulSoup(response.text, "html.parser")
#         doc = Document()

#         for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'span']):
#             text = tag.get_text(strip=True)
#             if not text:
#                 continue
#             if tag.name.startswith('h'):
#                 doc.add_heading(text, level=int(tag.name[-1]))
#             elif tag.name in ['p', 'span']:
#                 doc.add_paragraph(text)

#         page_name = urlparse(url).path.strip("/").replace("/", "_") or "index"
#         doc_file_path = os.path.join(domain_name, f"{page_name}.docx")
#         doc.save(doc_file_path)

#         print(f"Saved: {doc_file_path}")

#         with lock:
#             for link in soup.find_all("a", href=True):
#                 full_url = urljoin(url, link["href"])
#                 if (urlparse(full_url).netloc.replace("www.", "") == domain_name and
#                         full_url not in visited and
#                         not any(full_url.lower().endswith(ext) for ext in excluded_extensions)):
#                     pages_to_visit.add(full_url)

#     except Exception as e:
#         print(f"Failed to process {url}: {e}")

# def scrape_website(base_url):
#     headers = {
#         'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
#     }
#     excluded_extensions = (
#         '.mp3', '.mp4', '.wav', '.ogg', '.flac', '.aac',
#         '.avi', '.mov', '.wmv', '.m4a', '.webm', '.3gp',
#         '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg',
#         '.pdf', '.docx'
#     )

#     visited = set()
#     pages_to_visit = set([base_url])
#     domain_name = urlparse(base_url).netloc.replace("www.", "")
#     create_directory(domain_name)

#     print(f"Starting scrape for domain: {domain_name}")

#     while pages_to_visit:
#         current_batch = list(pages_to_visit)[:200]
#         pages_to_visit.difference_update(current_batch)

#         with ThreadPoolExecutor(max_workers=50) as executor:
#             futures = [
#                 executor.submit(
#                     scrape_single_page,
#                     url, domain_name, visited, pages_to_visit, excluded_extensions, headers
#                 ) for url in current_batch
#             ]
#             for future in as_completed(futures):
#                 future.result()

#     print(f"Scraping completed. Total pages visited: {len(visited)}")

def is_image_url(url):
    try:
        response = requests.head(url, timeout=5, allow_redirects=True)
        content_type = response.headers.get('Content-Type', '')
        return content_type.startswith('image/')
    except:
        return False

#aync function with multithreading to scrape the website
def scrape_single_page(url, domain_name, visited, pages_to_visit, excluded_extensions, headers):
    with lock:
        if url in visited or any(url.lower().endswith(ext) for ext in excluded_extensions) or is_image_url(url):
            return
        visited.add(url)

    logger.info(f"Visiting: {url}")
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")
        doc = Document()

        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'span']):
            text = tag.get_text(strip=True)
            if not text:
                continue
            if tag.name.startswith('h'):
                doc.add_heading(text, level=int(tag.name[-1]))
            elif tag.name in ['p', 'span']:
                doc.add_paragraph(text)

        page_name = urlparse(url).path.strip("/").replace("/", "_") or "index"
        doc_file_path = os.path.join(domain_name, f"{page_name}.docx")
        doc.save(doc_file_path)

        logger.info(f"Saved: {doc_file_path}")

        with lock:
            for link in soup.find_all("a", href=True):
                full_url = urljoin(url, link["href"])
                if (urlparse(full_url).netloc.replace("www.", "") == domain_name and
                        full_url not in visited and
                        not any(full_url.lower().endswith(ext) for ext in excluded_extensions)):
                    pages_to_visit.add(full_url)

    except Exception as e:
        logger.error(f"Failed to process {url}: {e}")

async def scrape_website(base_url):
    logger.info(f"Starting scrape for base URL: {base_url}")
    loop = asyncio.get_event_loop()
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
    }

    excluded_extensions = (
        '.mp3', '.mp4', '.wav', '.ogg', '.flac', '.aac',
        '.avi', '.mov', '.wmv', '.m4a', '.webm', '.3gp',
        '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg',
        '.pdf', '.docx'
    )

    visited = set()
    pages_to_visit = set([base_url])
    domain_name = urlparse(base_url).netloc.replace("www.", "")
    create_directory(domain_name)

    print(f"Starting scrape for domain: {domain_name}")

    while pages_to_visit:
        current_batch = list(pages_to_visit)[:300]
        pages_to_visit.difference_update(current_batch)

        # Run batch in ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=100) as executor:
            tasks = [
                loop.run_in_executor(
                    executor,
                    scrape_single_page,
                    url,
                    domain_name,
                    visited,
                    pages_to_visit,
                    excluded_extensions,
                    headers
                )
                for url in current_batch
            ]
            await asyncio.gather(*tasks)

    logger.info(f"Scraping completed. Total pages visited: {len(visited)}")

# Function to scrape the website
def scrape_website_old(base_url):
    visited = set()
    pages_to_visit = [base_url]

    # Create a directory to store the scraped pages
    domain_name = urlparse(base_url).netloc.replace("www.", "")
    create_directory(domain_name)

    print(pages_to_visit)
    while pages_to_visit:
        url = pages_to_visit.pop(0)

        if url in visited:
            continue

        visited.add(url)

        try:
            response = requests.get(url)
            response.raise_for_status()  # Check for HTTP errors

            # Parse the HTML content
            soup = BeautifulSoup(response.text, "html.parser")

            # Create a new Document
            doc = Document()

            # Scrape and add h, p, and span tags to the document
            for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'span']):
                if tag.name.startswith('h'):
                    doc.add_heading(tag.get_text(strip=True), level=int(tag.name[-1]))
                elif tag.name == 'p':
                    doc.add_paragraph(tag.get_text(strip=True))
                elif tag.name == 'span':
                    doc.add_paragraph(tag.get_text(strip=True), style='Normal')

            # Save the document as a .docx file
            page_name = urlparse(url).path.strip("/").replace("/", "_") or "index"
            doc_file_path = os.path.join(domain_name, f"{page_name}.docx")

            doc.save(doc_file_path)
            print(f"Saved: {doc_file_path}")

            # Find all links on the page
            for link in soup.find_all("a"):
                href = link.get("href")
                if href:
                    full_url = urljoin(url, href)
                    # Add the link if it is part of the same domain and not visited
                    if urlparse(full_url).netloc == domain_name and full_url not in visited:
                        pages_to_visit.append(full_url)

        except Exception as e:
            print(f"Failed to process {url}: {e}")

#Convert DOCX to PDF
# def convert_docx_to_pdf(docx_file, pdf_file):
#     # Read the docx file
#     doc = Document(docx_file)
#     text = []

#     # Extract text from docx
#     for para in doc.paragraphs:
#         text.append(para.text)

#     # Create a PDF with reportlab
#     c = canvas.Canvas(pdf_file, pagesize=letter)
#     width, height = letter
#     y = height - 40  # Start from the top

#     for line in text:
#         if y < 40:  # Check if we need a new page
#             c.showPage()
#             y = height - 40
#         c.drawString(40, y, line)
#         y -= 12  # Move down for the next line

#     c.save()
# Convert DOCX to PDF
def convert_docx_to_pdf2(docx_file, pdf_file):
    # Read the docx file
    doc = Document(docx_file)

    # Create a PDF with reportlab
    c = canvas.Canvas(pdf_file, pagesize=letter)
    width, height = letter
    y = height - 40  # Start from the top

    for para in doc.paragraphs:
        if not para.text.strip():  # Skip empty paragraphs
            continue

        # Determine font style and size based on the paragraph style
        if para.style.name == 'Heading 1':
            c.setFont("Helvetica-Bold", 14)
        elif para.style.name == 'Heading 2':
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 10)

        # Check if we need a new page
        if y < 40:  
            c.showPage()
            y = height - 40

        # Write the paragraph text
        c.drawString(40, y, para.text)
        y -= 14  # Move down for the next line, adjusting spacing as needed

    c.save()

def convert_docx_to_pdf(docx_file, pdf_file):
    # Read the docx file
    doc = Document(docx_file)

    # Create a PDF with reportlab
    c = canvas.Canvas(pdf_file, pagesize=letter)
    width, height = letter
    y = height - 40  # Start from the top

    # Define a margin for text wrapping
    left_margin = 40
    right_margin = width - 40

    for para in doc.paragraphs:
        if not para.text.strip():  # Skip empty paragraphs
            continue

        # Set font style and size based on the paragraph style
        if para.style.name == 'Heading 1':
            c.setFont("Helvetica-Bold", 14)
        elif para.style.name == 'Heading 2':
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 10)

        # Check if we need a new page
        if y < 40:  
            c.showPage()
            y = height - 40

        # Wrap text to fit within the page's margins
        wrapped_text = textwrap.fill(para.text, width=90)  # Adjust width as needed
        
        # Write the paragraph text line by line
        for line in wrapped_text.split("\n"):
            if y < 40:
                c.showPage()  # Create new page if space runs out
                y = height - 40
            c.drawString(left_margin, y, line)
            y -= 14  # Adjust spacing between lines

    c.save()

# def convert_multiple_docx_to_pdf(docx_folder, output_folder):
#     # Create output folder if it doesn't exist
#     if not os.path.exists(output_folder):
#         os.makedirs(output_folder)

#     for filename in os.listdir(docx_folder):
#         if filename.endswith('.docx'):
#             docx_file = os.path.join(docx_folder, filename)
#             pdf_file = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}.pdf")
#             convert_docx_to_pdf(docx_file, pdf_file)
#             print(f"Converted {docx_file} to {pdf_file}")

async def convert_multiple_docx_to_pdf(docx_folder, output_folder):
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, _sync_convert_multiple_docx_to_pdf, docx_folder, output_folder)

def _sync_convert_multiple_docx_to_pdf(docx_folder, output_folder):
    # Ensure output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        logger.info(f"Created output folder: {output_folder}")

    for filename in os.listdir(docx_folder):
        if filename.endswith('.docx'):
            docx_file = os.path.join(docx_folder, filename)
            pdf_file = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}.pdf")
            try:
                convert_docx_to_pdf(docx_file, pdf_file)
                logger.info(f"Converted {docx_file} to {pdf_file}")
            except Exception as e:
                logger.error(f"Failed to convert {docx_file}: {e}")

# Create embeddings from the converted PDFs
# Initialize Bedrock client and embeddings
load_dotenv()
# Retrieve AWS credentials from environment variables
bedrock = boto3.client(service_name='bedrock-runtime',
                        region_name='us-east-1'
                        )
bedrock_embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v2:0", client=bedrock)


# def data_ingestion(pdf_folder):
#     loader = PyPDFDirectoryLoader(pdf_folder)  # Folder containing PDF files
#     documents = loader.load()  # Load documents from PDFs
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
#     docs = text_splitter.split_documents(documents)  # Split documents into chunks
#     return docs

# async def data_ingestion(pdf_folder):
#     loop = asyncio.get_event_loop()
#     return await loop.run_in_executor(None, _sync_data_ingestion, pdf_folder)

async def data_ingestion(docx_folder):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _sync_data_ingestion, docx_folder)

# def _sync_data_ingestion(pdf_folder):
#     logger.info(f"Starting data ingestion from folder: {pdf_folder}")
    
#     try:
#         loader = UnstructuredWordDocumentLoader(pdf_folder)  
#         documents = loader.load()
#         logger.info(f"Loaded {len(documents)} documents")

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
#         docs = text_splitter.split_documents(documents)
#         logger.info(f"Split into {len(docs)} document chunks")

#         return docs
#     except Exception as e:
#         logger.error(f"Failed during data ingestion from {pdf_folder}: {e}")
#         return []

def _sync_data_ingestion(docx_folder):
    logger.info(f"Starting DOCX data ingestion from folder: {docx_folder}")
    documents = []

    try:
        for file in os.listdir(docx_folder):
            if file.endswith('.docx'):
                file_path = os.path.join(docx_folder, file)
                try:
                    loader = UnstructuredWordDocumentLoader(file_path)
                    docs = loader.load()
                    documents.extend(docs)
                    logger.info(f"Loaded {len(docs)} pages from: {file}")
                except Exception as e:
                    logger.warning(f"Failed to load DOCX {file}: {e}")

        if not documents:
            logger.warning("No DOCX documents found for ingestion.")
            return []

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
        docs = text_splitter.split_documents(documents)
        logger.info(f"Split into {len(docs)} document chunks")

        return docs

    except Exception as e:
        logger.error(f"Data ingestion failed for folder {docx_folder}: {e}")
        return []

# def get_vector_store(docs,domain_name):
#     dname = domain_name.replace(".","_")
#     vector_store_faiss = FAISS.from_documents(docs, bedrock_embeddings)  # Create vector store using embeddings
#     #FAISS.afrom_documents
#     vector_store_faiss.save_local("faiss_index_"+dname)  # Save vector store locally

async def get_vector_store(docs, domain_name):
    dname = domain_name.replace(".", "_")
    index_path = f"faiss_index_{dname}"

    logger.info(f"Starting vector store creation for domain: {domain_name}")
    logger.info(f"Embedding {len(docs)} document chunks")

    try:
        vector_store_faiss = await FAISS.afrom_documents(docs, bedrock_embeddings)  # Async embedding
        vector_store_faiss.save_local(index_path)
        logger.info(f"FAISS index saved at: {index_path}")
        return vector_store_faiss
    except Exception as e:
        logger.error(f"Failed to create vector store for {domain_name}: {e}")
        return None

# def onboard_business(business_name,business_website,business_desp):
#     website_url = business_website  # Replace with the target URL
#     folder_name = business_name 

#     # Create the directory if it doesn't exist
#     os.makedirs(folder_name, exist_ok=True)

#     #urls = extract_urls_to_csv(website_url, folder_name)
#     #extract_content_from_urls(urls, folder_name)
#     scrape_website(website_url)
#     domain_name = urlparse(website_url).netloc.replace("www.", "")
#     pdf_folder = os.path.join(domain_name, 'pdf')
#     print(pdf_folder)

#     print("CONVERTING DOCX to PDFS ---")
#     convert_multiple_docx_to_pdf(domain_name, pdf_folder)  # Convert DOCX to PDF

#     # Step 2: Ingest the converted PDFs and create embeddings
#     print("DATA INGESTION STEP IN PROGRESS")
#     docs = data_ingestion(domain_name)  # Ingest data from converted PDFs
#     print("GENERATING FAISS AND PICKLE")
#     get_vector_store(docs,domain_name)   # Create embeddings and save to FAISS index
#     print("PROCESS COMPLETED")
#     dname = domain_name.replace(".","_")
#     return "http://incognitochat.s3-website-us-east-1.amazonaws.com/?id="+dname

# Asynchronous function of onboard_business
async def onboard_business(business_name, business_website, business_desp):
    website_url = business_website
    folder_name = business_name

    # Create the directory if it doesn't exist
    os.makedirs(folder_name, exist_ok=True)
    logger.info(f"Created/confirmed folder: {folder_name}")

    #Scrape website
    logger.info(f"Starting scraping for: {website_url}")
    await scrape_website(website_url)

    domain_name = urlparse(website_url).netloc.replace("www.", "")
    # pdf_folder = os.path.join(domain_name, 'pdf')
    # logger.info(f"PDF output folder set to: {pdf_folder}")

    # #Convert DOCX to PDF
    # logger.info("Converting DOCX files to PDFs...")
    # await convert_multiple_docx_to_pdf(domain_name, pdf_folder)

    #Data ingestion
    logger.info("Starting data ingestion...")
    docs = await data_ingestion(domain_name)

    if not docs:
        logger.error("No documents found to embed. Check scraping and PDF conversion steps.")
        raise ValueError("No documents found to embed.")
    
    logger.info(f"Number of documents to embed: {len(docs)}")

    #Generate FAISS vector store
    logger.info("Generating FAISS index and saving locally...")
    await get_vector_store(docs, domain_name)

    logger.info("Business onboarding process completed successfully.")
    dname = domain_name.replace(".", "_")
    final_url = f"http://incognitochat.s3-website-us-east-1.amazonaws.com/?id={dname}"
    logger.info(f"Returning access URL: {final_url}")
    
    return final_url


def makepickel(business_website):
    domain_name = urlparse(business_website).netloc.replace("www.", "")
    pdf_folder = os.path.join(domain_name, 'pdf')
    print(pdf_folder)

    print("CONVERTING DOCX to PDFS ---")
    convert_multiple_docx_to_pdf(domain_name, pdf_folder)  # Convert DOCX to PDF

    # Step 2: Ingest the converted PDFs and create embeddings
    print("DATA INGESTION STEP IN PROGRESS")
    docs = data_ingestion(pdf_folder)  # Ingest data from converted PDFs
    print("GENERATING FAISS AND PICKLE")
    get_vector_store(docs,domain_name)   # Create embeddings and save to FAISS index
    print("PROCESS COMPLETED")
    dname = domain_name.replace(".","_")
    return "http://incognitochat.s3-website-us-east-1.amazonaws.com/?id="+dname






